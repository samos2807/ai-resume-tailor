"""Resume builder — v4. Matches reference template exactly.

Layout (from data/cv-template-reference.docx):
- Header: centered name (22pt bold navy), title (18pt bold navy), contact (10pt, one line)
- Section headers: 10.5pt bold navy, left-aligned, dotted separator above
- Project bullets: checkmark (12pt bold) + text (10pt)
- Military bullets: bold dot + text (10pt)
- Languages: separate section at bottom
"""

import json
import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

from matcher.cv_rules import (
    NAVY_BLUE, BLACK, LINK_GREEN, SEPARATOR_GRAY,
    FONT, NAME_PT, TITLE_PT, CONTACT_PT, HEADER_PT, BODY_PT,
    CHECKMARK_PT, BULLET_PT,
    MARGIN_CM, LINE_SPACING_PT,
    PHONE, EMAIL, LINKEDIN_URL, LINKEDIN_TEXT, GITHUB_URL, GITHUB_TEXT,
    SKILL_POOL, PROJECTS, EDUCATION, COURSEWORK_MAP,
    MILITARY, MILITARY_BULLET1, MILITARY_BULLET2,
    VOLUNTEERING_LINE, HOBBIES_LINE, LANGUAGES_LINE,
    BANNED_WORDS, COMPANY_DOMAINS,
    get_cv_generation_prompt,
)

logger = logging.getLogger(__name__)
OUTPUT_DIR = Path(__file__).parent.parent / "output" / "resumes"

# ── Colors (docx-ready) ──
BLUE = RGBColor(*NAVY_BLUE)
BLK = RGBColor(*BLACK)
LINK_CLR = RGBColor(*LINK_GREEN)
FONT_NAME = FONT

# ── Sizes ──
NAME_SIZE = Pt(NAME_PT)
TITLE_SIZE = Pt(TITLE_PT)
CONTACT_SIZE = Pt(CONTACT_PT)
SECTION_HEADER_SIZE = Pt(HEADER_PT)
BODY_SIZE = Pt(BODY_PT)
CHECKMARK_SIZE = Pt(CHECKMARK_PT)
BULLET_SIZE = Pt(BULLET_PT)
EDUCATION_SIZE = Pt(BODY_PT)

# ── Spacing ──
SECTION_AFTER = Pt(6)
BULLET_AFTER = Pt(1)
PROJECT_TITLE_BEFORE = Pt(10)
LINE_SPACING = Pt(LINE_SPACING_PT)
MARGIN_TOP = Cm(MARGIN_CM["top"])
MARGIN_BOTTOM = Cm(MARGIN_CM["bottom"])
MARGIN_LEFT = Cm(MARGIN_CM["left"])
MARGIN_RIGHT = Cm(MARGIN_CM["right"])


# ══════════════════════════════════════════════════════════════════════
# AI INTEGRATION
# ══════════════════════════════════════════════════════════════════════

def _call_ai_for_cv(job_title, company, description):
    try:
        import anthropic
        from config import ANTHROPIC_API_KEY
        if not ANTHROPIC_API_KEY:
            raise ValueError("No API key")
        client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
        prompt = get_cv_generation_prompt(job_title, company, description)
        response = client.messages.create(
            model="claude-sonnet-4-20250514", max_tokens=1000,
            messages=[{"role": "user", "content": prompt}],
        )
        text = response.content[0].text.strip()
        if text.startswith("```"):
            text = re.sub(r'^```\w*\n?', '', text)
            text = re.sub(r'\n?```$', '', text)
        return json.loads(text)
    except Exception as e:
        logger.warning("AI CV generation failed (%s), using fallback", e)
        return _smart_fallback(job_title, company, description)


def _smart_fallback(job_title, company, description=""):
    text = f"{job_title} {description}".lower()
    jt = _detect_job_type(text)
    return {
        "subtitle": _fallback_subtitle(jt, job_title),
        "about": _fallback_about(jt, company),
        "skills_custom": None,
        "skills": _fallback_skills(text),
        "project_order": _fallback_project_order(jt),
        "project_expansion": {_fallback_project_order(jt)[0]: "full",
                              _fallback_project_order(jt)[1]: "medium",
                              _fallback_project_order(jt)[2]: "compressed"},
        "military_tone": "default",
        "coursework": jt if jt in COURSEWORK_MAP else "general",
    }


def _detect_job_type(text):
    if any(w in text for w in ["verification", "dv ", "formal", "dft"]): return "verification"
    if any(w in text for w in ["physical design", "layout", "p&r", "floorplan"]): return "physical_design"
    if any(w in text for w in ["fpga", "embedded", "firmware"]): return "fpga_embedded"
    if any(w in text for w in ["asic", "rtl", "digital design", "synthesis"]): return "asic_rtl"
    if any(w in text for w in ["application engineer", "fae", "customer"]): return "application"
    if any(w in text for w in ["analog", "mixed signal"]): return "analog"
    return "general"


def _fallback_subtitle(jt, title):
    t = title.lower()
    if "application" in t: return "Junior Application Engineer"
    if "verification" in t: return "Design Verification Engineer"
    if "physical design" in t: return "Physical Design Engineer"
    if "fpga" in t: return "FPGA Design Engineer"
    if "embedded" in t: return "Embedded Systems Engineer"
    if "asic" in t or "rtl" in t: return "RTL & ASIC Design Engineer"
    return {"verification": "Design Verification Engineer", "physical_design": "Physical Design Engineer",
            "fpga_embedded": "FPGA & Embedded Engineer", "asic_rtl": "RTL & ASIC Design Engineer",
            "application": "Junior Application Engineer"}.get(jt, "Hardware Design Engineer")


def _fallback_skills(text):
    scores = {}
    for key, cat in SKILL_POOL.items():
        if key == "scripting": continue
        scores[key] = sum(1 for t in cat["triggers"] if t in text)
    ranked = sorted(scores.keys(), key=lambda k: scores[k], reverse=True)
    top = [k for k in ranked if scores[k] > 0][:2]
    if not top: top = ranked[:2]
    if any(w in text for w in ["python", "tcl", "scripting"]): top.append("scripting")
    return top


def _fallback_project_order(jt):
    return {"verification": ["riscv", "vlsi_lab", "fatigue"], "asic_rtl": ["riscv", "vlsi_lab", "fatigue"],
            "fpga_embedded": ["fatigue", "riscv", "vlsi_lab"], "physical_design": ["vlsi_lab", "riscv", "fatigue"],
            "analog": ["vlsi_lab", "riscv", "fatigue"], "application": ["riscv", "vlsi_lab", "fatigue"],
            "embedded": ["fatigue", "microblaze", "riscv"]}.get(jt, ["riscv", "fatigue", "vlsi_lab"])


def _fallback_about(jt, company):
    domain = COMPANY_DOMAINS.get(company.lower().strip(), "")
    b, be = "{bold}", "{/bold}"
    cl = company.lower().strip()
    if cl in ("intel", "intel corporation"):
        closer = f" Eager to apply strong RTL foundations and systematic validation mindset to {b}{company}{be}'s silicon development."
    elif cl in ("nvidia", "mellanox"):
        closer = f" Eager to apply strong design foundations and fast EDA tool learning to {b}{company}{be}'s next-generation silicon."
    elif domain:
        closer = f" Eager to apply strong design foundations to {b}{company}{be}'s {domain}."
    else:
        closer = f" Eager to contribute to the engineering team at {b}{company}{be}."
    intros = {
        "verification": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}RTL design{be}, {b}verification{be}, {b}ASIC implementation{be}, and {b}debug{be}. Experienced in taking designs from specification through RTL, {b}testbench development{be}, {b}functional verification{be}, and system validation. Worked with {b}Verilog{be}, {b}Cadence{be} Genus, Innovus, and {b}Xilinx Vivado{be}.",
        "asic_rtl": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}RTL design{be}, {b}ASIC implementation{be}, and {b}digital design{be}. Experienced in taking designs from specification through {b}RTL{be}, {b}synthesis{be}, {b}physical layout{be}, {b}DRC/LVS{be}, and system validation. Worked with {b}Verilog{be}, {b}Cadence Genus{be}, {b}Innovus{be}, and {b}Xilinx Vivado{be}.",
        "physical_design": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}physical design{be}, {b}ASIC implementation{be}, and analog VLSI. Experienced in taking designs from RTL through {b}synthesis{be}, {b}floorplanning{be}, {b}place-and-route{be}, {b}STA{be}, and {b}DRC/LVS{be} signoff. Worked with {b}Cadence Genus{be}, {b}Innovus{be}, {b}Virtuoso{be}, and {b}Xilinx Vivado{be}.",
        "fpga_embedded": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}FPGA development{be}, {b}HW/SW integration{be}, and {b}embedded systems{be}. Experienced in taking designs from specification through implementation, {b}board bring-up{be}, and real-time system validation. Worked with {b}Xilinx Vivado{be}, {b}Kria KR260{be}, {b}MicroBlaze{be}, and {b}Embedded C{be}.",
        "application": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}RTL design{be}, {b}FPGA development{be}, and {b}ASIC implementation{be}. Experienced in taking designs from specification through RTL, synthesis, physical layout, and system validation. Worked with {b}Cadence{be} Genus, Innovus, Virtuoso, and {b}Xilinx Vivado{be}.",
        "general": f"Electrical & Electronics Engineering graduate with hands-on experience in {b}RTL design{be}, {b}FPGA development{be}, {b}ASIC implementation{be}, and analog VLSI. Experienced in taking designs from specification through RTL, synthesis, physical layout, DRC/LVS, and system validation. Worked with {b}Cadence{be} Genus, Innovus, Virtuoso, and {b}Xilinx Vivado{be}.",
    }
    return intros.get(jt, intros["general"]) + closer


# ══════════════════════════════════════════════════════════════════════
# TEXT FORMATTING
# ══════════════════════════════════════════════════════════════════════

def _parse_bold_text(paragraph, text, font_size=None, font_color=None,
                     bold_keywords=None, base_bold=False):
    sz = font_size or BODY_SIZE
    color = font_color or BLK
    parts = re.split(r'\{bold\}|\{/bold\}', text)
    is_bold = False
    for part in parts:
        if part:
            if not is_bold and bold_keywords:
                for sub_text, sub_bold in _split_by_keywords(part, bold_keywords):
                    run = paragraph.add_run(sub_text)
                    run.bold = sub_bold or base_bold
                    run.font.size = sz
                    run.font.color.rgb = color
                    run.font.name = FONT_NAME
            else:
                run = paragraph.add_run(part)
                run.bold = is_bold or base_bold
                run.font.size = sz
                run.font.color.rgb = color
                run.font.name = FONT_NAME
        is_bold = not is_bold


def _split_by_keywords(text, keywords):
    if not keywords:
        return [(text, False)]
    sorted_kw = sorted(keywords, key=len, reverse=True)
    parts = []
    for kw in sorted_kw:
        if len(kw) < 3:
            continue
        escaped = re.escape(kw)
        if len(kw) <= 4 and kw.isalpha():
            parts.append(r'(?<![a-zA-Z0-9])' + escaped + r'(?![a-zA-Z0-9])')
        else:
            parts.append(escaped)
    pattern = '|'.join(parts)
    if not pattern:
        return [(text, False)]
    result = []
    last_end = 0
    for match in re.finditer(pattern, text, re.IGNORECASE):
        if match.start() > last_end:
            result.append((text[last_end:match.start()], False))
        result.append((text[match.start():match.end()], True))
        last_end = match.end()
    if last_end < len(text):
        result.append((text[last_end:], False))
    return result if result else [(text, False)]


# ══════════════════════════════════════════════════════════════════════
# KEYWORD EXTRACTION
# ══════════════════════════════════════════════════════════════════════

_RELATED_TERMS = {
    "sta_timing": ["STA", "Static Timing Analysis", "timing analysis", "signoff STA"],
    "par_place": ["P&R", "Place and Route", "Innovus"],
    "drc_lvs": ["DRC/LVS", "DRC", "LVS", "physical verification", "signoff"],
    "rtl_gds": ["RTL-to-GDS", "RTL2GDS", "ASIC flow", "implementation flow"],
    "synthesis": ["synthesis", "logic synthesis", "Genus"],
    "pex_parasitic": ["PEX", "parasitic extraction", "post-layout simulation", "Quantus PEX"],
    "fpga_tools": ["FPGA", "Vivado", "Xilinx", "Kria", "KR260", "SoC-FPGA"],
    "cadence_tools": ["Cadence", "Virtuoso", "Innovus", "Genus", "Spectre", "PVS"],
    "rtl_hdl": ["Verilog", "VHDL", "RTL", "RTL design", "RTL2GDS", "RTL-to-GDS"],
    "verification_dv": ["verification", "validation", "simulation", "gate-level", "corner-case"],
}

_CONTEXT_DOMAINS = {
    "verification": ["simulation", "gate-level", "corner cases", "debugging", "verification", "validation"],
    "physical design": ["layout", "DRC", "LVS", "PEX", "timing analysis", "Innovus", "Virtuoso", "signoff", "P&R", "RTL2GDS", "Cadence"],
    "fpga": ["FPGA", "Vivado", "Xilinx", "Kria", "KR260", "SoC", "Vitis AI", "DPU"],
    "asic": ["ASIC", "RTL", "RTL2GDS", "Genus", "Innovus", "synthesis", "P&R", "STA", "DRC", "Verilog", "Tcl", "Cadence"],
}


def _extract_job_keywords(title, description=""):
    text = f"{title} {description}".lower()
    if not text.strip():
        return []
    bold_set = set()
    for terms in _RELATED_TERMS.values():
        for term in terms:
            if term.lower() in text:
                for t in terms:
                    bold_set.add(t)
                break
    for trigger, expansion in _CONTEXT_DOMAINS.items():
        if trigger.lower() in text:
            for t in expansion:
                bold_set.add(t)
    if len(bold_set) < 5:
        for t in ["Cadence", "Genus", "Innovus", "Virtuoso", "Vivado", "Verilog", "Python", "Tcl"]:
            bold_set.add(t)
    return [kw for kw in bold_set if len(kw) >= 3]


# ══════════════════════════════════════════════════════════════════════
# HYPERLINK HELPER
# ══════════════════════════════════════════════════════════════════════

def _add_hyperlink(paragraph, url, text, font_size=None, color_hex="2B5797"):
    sz = font_size or CONTACT_SIZE
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = parse_xml(f'<w:hyperlink {nsdecls("w")} xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships" r:id="{r_id}"/>')
    new_run = parse_xml(
        f'<w:r {nsdecls("w")}>'
        f'  <w:rPr><w:rFonts w:ascii="{FONT_NAME}" w:hAnsi="{FONT_NAME}"/>'
        f'    <w:sz w:val="{int(sz.pt * 2)}"/><w:color w:val="{color_hex}"/></w:rPr>'
        f'  <w:t>{text}</w:t></w:r>')
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


# ══════════════════════════════════════════════════════════════════════
# DOCUMENT BUILDER
# ══════════════════════════════════════════════════════════════════════

def _add_dotted_separator(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = Pt(1)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'dotted')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), SEPARATOR_GRAY)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_section_header(doc, text):
    _add_dotted_separator(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = SECTION_AFTER
    p.paragraph_format.line_spacing = LINE_SPACING
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = SECTION_HEADER_SIZE
    run.font.color.rgb = BLUE
    run.font.name = FONT_NAME


def _add_footer_line(doc):
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(3)
    p_foot.paragraph_format.space_after = Pt(0)
    r_foot = p_foot.add_run("AI-tailored by a system I built — ")
    r_foot.italic = True
    r_foot.font.size = Pt(8)
    r_foot.font.name = FONT_NAME
    r_foot.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    _add_hyperlink(p_foot, "https://github.com/samos2807/ai-resume-tailor", "samos2807/ai-resume-tailor", font_size=Pt(8))


def build_cv(cv_data, keywords=None, footer_placement="bottom"):
    keywords = keywords or []
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = MARGIN_TOP
    section.bottom_margin = MARGIN_BOTTOM
    section.left_margin = MARGIN_LEFT
    section.right_margin = MARGIN_RIGHT

    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = BODY_SIZE
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = LINE_SPACING

    # ════════════ HEADER (centered) ════════════

    # Name — biggest, bold, centered, navy
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_name.paragraph_format.space_after = Pt(8)
    run = p_name.add_run("Amos Sarusi")
    run.bold = True
    run.font.size = NAME_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLUE

    # Job title — one size smaller, bold, centered, navy
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    run = p_title.add_run(cv_data.get("subtitle", "Hardware Design Engineer"))
    run.bold = True
    run.font.size = TITLE_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLUE

    # Contact — one line, centered
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_contact.paragraph_format.space_after = Pt(0)
    run = p_contact.add_run(f"{PHONE}  |  {EMAIL}  |  ")
    run.font.size = CONTACT_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK
    _add_hyperlink(p_contact, LINKEDIN_URL, LINKEDIN_TEXT, font_size=CONTACT_SIZE)
    run = p_contact.add_run("  |  ")
    run.font.size = CONTACT_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK
    _add_hyperlink(p_contact, GITHUB_URL, GITHUB_TEXT, font_size=CONTACT_SIZE)

    # ════════════ ABOUT ════════════
    _add_section_header(doc, "ABOUT")
    about_text = cv_data.get("about", "")
    if isinstance(about_text, list):
        about_text = " ".join(about_text)
    p_about = doc.add_paragraph()
    p_about.paragraph_format.space_after = Pt(1)
    p_about.paragraph_format.space_before = Pt(1)
    _parse_bold_text(p_about, about_text, font_size=BODY_SIZE, font_color=BLK, bold_keywords=keywords)

    # ════════════ TECHNICAL SKILLS ════════════
    _add_section_header(doc, "TECHNICAL SKILLS")
    skills_custom = cv_data.get("skills_custom")
    if skills_custom:
        skill_lines = skills_custom
    else:
        skill_keys = cv_data.get("skills", ["digital_rtl", "asic_flow"])
        skill_lines = [{"label": SKILL_POOL[k]["label"], "content": SKILL_POOL[k]["content"]}
                       for k in skill_keys if k in SKILL_POOL]

    for sl in skill_lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(f"{sl['label']}: ")
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.name = FONT_NAME
        run.font.color.rgb = BLK
        if keywords:
            for text, is_bold in _split_by_keywords(sl['content'], keywords):
                run = p.add_run(text)
                run.bold = is_bold
                run.font.size = BODY_SIZE
                run.font.name = FONT_NAME
                run.font.color.rgb = BLK
        else:
            run = p.add_run(sl['content'])
            run.font.size = BODY_SIZE
            run.font.name = FONT_NAME
            run.font.color.rgb = BLK

    # ════════════ PROJECTS (checkmarks ✓) ════════════
    _add_section_header(doc, "PROJECTS")
    projects_custom = cv_data.get("projects_custom")

    def _add_checkmark_bullet(doc, bullet_text, keywords):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = BULLET_AFTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        # Checkmark at 12pt bold
        run = p.add_run("\u2713 ")
        run.bold = True
        run.font.size = CHECKMARK_SIZE
        run.font.name = FONT_NAME
        run.font.color.rgb = BLK
        _parse_bold_text(p, bullet_text, font_size=BODY_SIZE, font_color=BLK, bold_keywords=keywords)

    def _add_project_header(doc, title, meta):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = PROJECT_TITLE_BEFORE
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(title)
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.name = FONT_NAME
        run.font.color.rgb = BLK
        if meta:
            run = p.add_run("  |  ")
            run.font.size = BODY_SIZE
            run.font.name = FONT_NAME
            run.font.color.rgb = BLK
            # Bold Cadence in meta
            if keywords:
                for text, is_bold in _split_by_keywords(meta, keywords):
                    run = p.add_run(text)
                    run.bold = is_bold
                    run.font.size = BODY_SIZE
                    run.font.name = FONT_NAME
                    run.font.color.rgb = BLK
            else:
                run = p.add_run(meta)
                run.font.size = BODY_SIZE
                run.font.name = FONT_NAME
                run.font.color.rgb = BLK

    if projects_custom:
        for idx, proj in enumerate(projects_custom):
            if idx == 0 and footer_placement == "under_first_project":
                _add_footer_line(doc)
            _add_project_header(doc, proj['title'], proj.get('meta', ''))
            for bullet in proj.get('bullets', []):
                _add_checkmark_bullet(doc, bullet, keywords)
    else:
        project_order = cv_data.get("project_order", ["riscv", "fatigue", "vlsi_lab"])
        project_expansion = cv_data.get("project_expansion", {})
        # Enforce: riscv and fatigue always full (3 bullets)
        for key in ("riscv", "fatigue"):
            if key in project_expansion:
                project_expansion[key] = "full"
        coursework_key = cv_data.get("coursework", "general")
        jt_map = {"verification": "verification", "asic_rtl": "asic_rtl",
                  "physical_design": "physical_design", "fpga_embedded": "default",
                  "analog": "physical_design"}
        bullet_type = jt_map.get(coursework_key, "default")
        for idx, proj_key in enumerate(project_order[:3]):
            proj = PROJECTS.get(proj_key)
            if not proj: continue
            expansion = project_expansion.get(proj_key, "full" if proj_key in ("riscv", "fatigue") else "medium")
            if idx == 0 and footer_placement == "under_first_project":
                _add_footer_line(doc)
            _add_project_header(doc, proj['title'], proj.get('meta', ''))
            if expansion == "compressed":
                _add_checkmark_bullet(doc, proj["bullet_compressed"], keywords)
            else:
                by_type = proj.get("bullets_by_type")
                if by_type:
                    bullets = list(by_type.get(bullet_type, by_type.get("default", [])))
                else:
                    bullets = list(proj.get("bullets_full") or [])
                if expansion == "full":
                    bullets = bullets[:3]
                elif expansion == "medium":
                    bullets = bullets[:2]
                for bullet in bullets:
                    _add_checkmark_bullet(doc, bullet, keywords)

    # ════════════ EDUCATION ════════════
    _add_section_header(doc, "EDUCATION")
    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_before = Pt(1)
    p_edu.paragraph_format.space_after = Pt(0)
    run = p_edu.add_run(EDUCATION['degree'])
    run.bold = True
    run.font.size = EDUCATION_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK
    run = p_edu.add_run(f",  {EDUCATION['university']}, {EDUCATION['years']}  |  {EDUCATION['gpa']}")
    run.font.size = EDUCATION_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK

    coursework_key = cv_data.get("coursework", "general")
    coursework = COURSEWORK_MAP.get(coursework_key, COURSEWORK_MAP["general"])
    p_cw = doc.add_paragraph()
    p_cw.paragraph_format.space_after = Pt(0)
    # Bold course names that match job keywords
    _parse_bold_text(p_cw, coursework, font_size=EDUCATION_SIZE, font_color=BLK, bold_keywords=keywords)

    # ════════════ MILITARY SERVICE (• bullets) ════════════
    _add_section_header(doc, "MILITARY SERVICE")
    p_mil = doc.add_paragraph()
    p_mil.paragraph_format.space_before = Pt(1)
    p_mil.paragraph_format.space_after = Pt(0)
    run = p_mil.add_run(MILITARY['title_bold'])
    run.bold = True
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK
    run = p_mil.add_run(MILITARY['title_regular'])
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK

    mil_tone = cv_data.get("military_tone", "default")
    mil_bullets = [
        MILITARY_BULLET1.get(mil_tone, MILITARY_BULLET1["default"]),
        MILITARY_BULLET2.get(mil_tone, MILITARY_BULLET2["default"]),
    ]
    for bullet_text in mil_bullets:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = BULLET_AFTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        p.paragraph_format.first_line_indent = Inches(-0.2)
        # Bold bullet point •
        run = p.add_run("\u2022 ")
        run.bold = True
        run.font.size = BODY_SIZE
        run.font.name = FONT_NAME
        run.font.color.rgb = BLK
        _parse_bold_text(p, bullet_text, font_size=BODY_SIZE, font_color=BLK, bold_keywords=keywords)

    # ════════════ VOLUNTEERING & HOBBIES ════════════
    _add_section_header(doc, "VOLUNTEERING & HOBBIES")
    p_vol = doc.add_paragraph()
    p_vol.paragraph_format.space_before = Pt(1)
    p_vol.paragraph_format.space_after = Pt(0)
    run = p_vol.add_run(VOLUNTEERING_LINE)
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK

    p_hob = doc.add_paragraph()
    p_hob.paragraph_format.space_after = Pt(0)
    run = p_hob.add_run(HOBBIES_LINE)
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK

    # ════════════ LANGUAGES ════════════
    _add_section_header(doc, "LANGUAGES")
    p_lang = doc.add_paragraph()
    p_lang.paragraph_format.space_before = Pt(1)
    p_lang.paragraph_format.space_after = Pt(0)
    run = p_lang.add_run(LANGUAGES_LINE)
    run.font.size = BODY_SIZE
    run.font.name = FONT_NAME
    run.font.color.rgb = BLK

    # ════════════ FOOTER — AI-tailored note (bottom placement) ════════════
    if footer_placement == "bottom":
        _add_footer_line(doc)

    # ════════════ POST-PROCESS BOLD ════════════
    if keywords:
        _post_process_bold(doc, keywords)

    return doc


def _post_process_bold(doc, keywords):
    if not keywords: return
    sorted_kw = sorted([kw for kw in keywords if len(kw) >= 3], key=len, reverse=True)
    if not sorted_kw: return
    parts = []
    for kw in sorted_kw:
        escaped = re.escape(kw)
        if len(kw) <= 4 and kw.isalpha():
            parts.append(r'(?<![a-zA-Z0-9])' + escaped + r'(?![a-zA-Z0-9])')
        else:
            parts.append(escaped)
    pattern = re.compile('(' + '|'.join(parts) + ')', re.IGNORECASE)

    for paragraph in doc.paragraphs:
        runs = paragraph.runs
        if runs and runs[0].font.color.rgb == BLUE: continue
        full_text = paragraph.text
        if not full_text or len(full_text) < 5: continue
        already_bold = sum(len(r.text) for r in runs if r.bold and r.text)
        total = len(full_text)
        if total > 0 and already_bold / total > 0.30: continue

        new_runs = []
        needs_rebuild = False
        new_bold = 0
        for run in runs:
            text = run.text
            if not text or run.bold:
                new_runs.append({'text': text, 'bold': run.bold, 'italic': run.italic,
                                'size': run.font.size, 'color': run.font.color.rgb, 'name': run.font.name})
                continue
            if not pattern.search(text):
                new_runs.append({'text': text, 'bold': run.bold, 'italic': run.italic,
                                'size': run.font.size, 'color': run.font.color.rgb, 'name': run.font.name})
                continue
            needs_rebuild = True
            for part in pattern.split(text):
                if not part: continue
                is_match = bool(pattern.fullmatch(part))
                if is_match:
                    if (already_bold + new_bold + len(part)) / total > 0.35:
                        is_match = False
                    else:
                        new_bold += len(part)
                new_runs.append({'text': part, 'bold': is_match or run.bold, 'italic': run.italic,
                                'size': run.font.size, 'color': run.font.color.rgb, 'name': run.font.name})

        if needs_rebuild:
            for r in list(paragraph._p.findall(qn('w:r'))):
                paragraph._p.remove(r)
            for rd in new_runs:
                if not rd['text']: continue
                r = paragraph.add_run(rd['text'])
                r.bold = rd['bold']
                r.italic = rd.get('italic', False)
                if rd.get('size'): r.font.size = rd['size']
                if rd.get('color'): r.font.color.rgb = rd['color']
                if rd.get('name'): r.font.name = rd['name']


# ══════════════════════════════════════════════════════════════════════
# TITLE CLEANING
# ══════════════════════════════════════════════════════════════════════

def clean_title(raw_title):
    if not raw_title: return "Hardware Engineer"
    title = raw_title.strip()
    title = re.sub(r"\s*Apply\s+.*$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"\s*VIEW POSITION\s*$", "", title, flags=re.IGNORECASE)
    title = re.sub(r"[\n\r]+", " ", title)
    title = re.sub(r"\s{2,}", " ", title)
    title = title.strip(" .,;|\u00b7\u2022\t")
    return title if title else "Hardware Engineer"


def sanitize_filename(text):
    return re.sub(r'[^\w\-]', '-', text.lower()).strip('-')[:40]


# ══════════════════════════════════════════════════════════════════════
# PUBLIC API
# ══════════════════════════════════════════════════════════════════════

def generate_cv_for_job(job):
    raw_title = job.get("title", "Hardware Engineer")
    company = job.get("company", "Unknown")
    description = job.get("description", "") or ""
    title = clean_title(raw_title)
    cv_data = _call_ai_for_cv(title, company, description)
    keywords = _extract_job_keywords(title, description)
    doc = build_cv(cv_data, keywords=keywords)

    company_folder = OUTPUT_DIR / sanitize_filename(company)
    company_folder.mkdir(parents=True, exist_ok=True)
    filename = f"resume-{sanitize_filename(company)}-{sanitize_filename(title)}"
    docx_path = company_folder / f"{filename}.docx"
    pdf_path = company_folder / f"{filename}.pdf"

    doc.save(str(docx_path))
    logger.info("Saved DOCX: %s", docx_path)
    try:
        import docx2pdf, time
        try: docx2pdf.convert(str(docx_path), str(pdf_path))
        except: time.sleep(1); docx2pdf.convert(str(docx_path), str(pdf_path))
        logger.info("Saved PDF: %s", pdf_path)
    except Exception as e:
        logger.warning("PDF conversion failed: %s", e)
        pdf_path = None
    return docx_path, pdf_path


def generate_cv_direct(job_title, company, description="", cv_data_override=None):
    title = clean_title(job_title)
    cv_data = cv_data_override if cv_data_override else _call_ai_for_cv(title, company, description)
    keywords = _extract_job_keywords(title, description)
    doc = build_cv(cv_data, keywords=keywords)

    company_folder = OUTPUT_DIR / sanitize_filename(company)
    company_folder.mkdir(parents=True, exist_ok=True)
    filename = f"resume-{sanitize_filename(company)}-{sanitize_filename(title)}"
    docx_path = company_folder / f"{filename}.docx"
    pdf_path = company_folder / f"{filename}.pdf"

    doc.save(str(docx_path))
    logger.info("Saved DOCX: %s", docx_path)
    try:
        import docx2pdf, time
        try: docx2pdf.convert(str(docx_path), str(pdf_path))
        except: time.sleep(1); docx2pdf.convert(str(docx_path), str(pdf_path))
        logger.info("Saved PDF: %s", pdf_path)
    except Exception as e:
        logger.warning("PDF conversion failed: %s", e)
        pdf_path = None
    return docx_path, pdf_path
